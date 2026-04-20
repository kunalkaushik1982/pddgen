r"""
Generate docs/templates/flowlens-sop-template.docx for docxtpl SOP export tests.

When ``sop.multi_process`` is true, the procedure body repeats once per
``sop.process_document_blocks`` item; the cover block stays single.

Run from repo root: python docs/templates/build_flowlens_sop_template.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document as open_document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def p(document: DocxDocument, text: str, *, style: str | None = None) -> None:
    if style:
        document.add_paragraph(text, style=style)
    else:
        document.add_paragraph(text)


def _add_sop_body_single(doc: DocxDocument) -> None:
    """Full SOP body using top-level ``sop.*`` (single workflow area or combined view)."""
    p(doc, "Purpose", style="Heading 1")
    p(doc, "{{ sop.purpose }}")

    p(doc, "")
    p(doc, "Scope", style="Heading 1")
    p(doc, "{{ sop.scope }}")

    p(doc, "")
    p(doc, "Applications and systems", style="Heading 1")
    p(doc, "{% for app in sop.applications %}")
    p(doc, "{{ app }}", style="List Bullet")
    p(doc, "{% endfor %}")
    p(doc, "{% if not sop.applications %}")
    p(doc, "No applications were recorded in the session evidence.", style="Intense Quote")
    p(doc, "{% endif %}")

    p(doc, "")
    p(doc, "Prerequisites", style="Heading 1")
    p(doc, "{% for item in sop.prerequisites %}")
    p(doc, "{{ item }}", style="List Bullet")
    p(doc, "{% endfor %}")

    p(doc, "")
    p(doc, "Roles and responsibilities", style="Heading 1")
    p(doc, "{% for row in sop.responsibilities %}")
    p(doc, "{{ row.role }}: {{ row.responsibility }}")
    p(doc, "{% endfor %}")

    p(doc, "")
    p(doc, "Control points and validation", style="Heading 1")
    p(doc, "{% for c in sop.controls %}")
    p(doc, "{{ c }}", style="List Bullet")
    p(doc, "{% endfor %}")

    p(doc, "")
    p(doc, "Expected outcomes", style="Heading 1")
    p(doc, "{% for e in sop.expected_outcomes %}")
    p(doc, "{{ e }}", style="List Bullet")
    p(doc, "{% endfor %}")

    p(doc, "")
    p(doc, "Procedure", style="Heading 1")
    p(doc, "{% for sec in sop.procedure_sections %}")
    p(doc, "{{ sec.title }}", style="Heading 2")
    p(doc, "{{ sec.summary }}")
    p(doc, "Objective: {{ sec.objective }}")
    p(doc, "{% for step in sec.steps %}")
    p(doc, "{{ step.step_number }}. {{ step.instruction }} ({{ step.system }})")
    p(doc, "{% if step.primary_screenshot_image %}")
    p(doc, "{{ step.primary_screenshot_image }}")
    p(doc, "{% endif %}")
    p(doc, "{% endfor %}")
    p(doc, "{% if sec.diagram_image %}")
    p(doc, "{{ sec.diagram_image }}")
    p(doc, "{% endif %}")
    p(doc, "{% endfor %}")

    p(doc, "")
    p(doc, "Supporting notes", style="Heading 1")
    p(doc, "{% for note in sop.supporting_notes %}")
    p(doc, "{{ note.text }}", style="List Bullet")
    p(doc, "{% endfor %}")

    p(doc, "")
    p(doc, "Evidence summary", style="Heading 1")
    p(
        doc,
        "Sections: {{ sop.evidence_summary.workflow_sections }} · "
        "Steps: {{ sop.evidence_summary.procedural_steps }} · "
        "Notes: {{ sop.evidence_summary.supporting_notes }} · "
        "Screenshots: {{ sop.evidence_summary.primary_screenshots }}",
    )


def _add_sop_body_per_block(doc: DocxDocument) -> None:
    """Same structure as single-process, scoped to ``block`` inside the outer loop."""
    p(doc, "{{ block.process_title }}", style="Heading 2")

    p(doc, "Purpose", style="Heading 1")
    p(doc, "{{ block.purpose }}")

    p(doc, "")
    p(doc, "Scope", style="Heading 1")
    p(doc, "{{ block.scope }}")

    p(doc, "")
    p(doc, "Applications and systems", style="Heading 1")
    p(doc, "{% for app in block.applications %}")
    p(doc, "{{ app }}", style="List Bullet")
    p(doc, "{% endfor %}")
    p(doc, "{% if not block.applications %}")
    p(doc, "No applications were recorded for this workflow section.", style="Intense Quote")
    p(doc, "{% endif %}")

    p(doc, "")
    p(doc, "Prerequisites", style="Heading 1")
    p(doc, "{% for item in block.prerequisites %}")
    p(doc, "{{ item }}", style="List Bullet")
    p(doc, "{% endfor %}")

    p(doc, "")
    p(doc, "Roles and responsibilities", style="Heading 1")
    p(doc, "{% for row in block.responsibilities %}")
    p(doc, "{{ row.role }}: {{ row.responsibility }}")
    p(doc, "{% endfor %}")

    p(doc, "")
    p(doc, "Control points and validation", style="Heading 1")
    p(doc, "{% for c in block.controls %}")
    p(doc, "{{ c }}", style="List Bullet")
    p(doc, "{% endfor %}")

    p(doc, "")
    p(doc, "Expected outcomes", style="Heading 1")
    p(doc, "{% for e in block.expected_outcomes %}")
    p(doc, "{{ e }}", style="List Bullet")
    p(doc, "{% endfor %}")

    p(doc, "")
    p(doc, "Procedure", style="Heading 1")
    p(doc, "{% for sec in block.procedure_sections %}")
    p(doc, "{{ sec.title }}", style="Heading 2")
    p(doc, "{{ sec.summary }}")
    p(doc, "Objective: {{ sec.objective }}")
    p(doc, "{% for step in sec.steps %}")
    p(doc, "{{ step.step_number }}. {{ step.instruction }} ({{ step.system }})")
    p(doc, "{% if step.primary_screenshot_image %}")
    p(doc, "{{ step.primary_screenshot_image }}")
    p(doc, "{% endif %}")
    p(doc, "{% endfor %}")
    p(doc, "{% if sec.diagram_image %}")
    p(doc, "{{ sec.diagram_image }}")
    p(doc, "{% endif %}")
    p(doc, "{% endfor %}")

    p(doc, "")
    p(doc, "Supporting notes", style="Heading 1")
    p(doc, "{% for note in block.supporting_notes %}")
    p(doc, "{{ note.text }}", style="List Bullet")
    p(doc, "{% endfor %}")

    p(doc, "")
    p(doc, "Evidence summary", style="Heading 1")
    p(
        doc,
        "Sections: {{ block.evidence_summary.workflow_sections }} · "
        "Steps: {{ block.evidence_summary.procedural_steps }} · "
        "Notes: {{ block.evidence_summary.supporting_notes }} · "
        "Screenshots: {{ block.evidence_summary.primary_screenshots }}",
    )


def main() -> None:
    out = Path(__file__).resolve().parent / "flowlens-sop-template.docx"
    doc = open_document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FlowLens Standard Operating Procedure")
    run.bold = True
    run.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("{{ sop.title }}").italic = True

    p(doc, "")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "Owner: {{ sop.owner_id }}  |  Session: {{ sop.session_id }}  |  Status: {{ sop.status }}\n"
        "Generated: {{ sop.generated_at }}  |  Diagram type: {{ sop.diagram_type }}"
    )

    p(doc, "")
    p(
        doc,
        "This document captures the operating procedure from the reviewed walkthrough. "
        "Steps, controls, and screenshots reflect session evidence where available.",
        style="Intense Quote",
    )

    p(doc, "{% if sop.multi_process %}")
    p(doc, "{% for block in sop.process_document_blocks %}")
    _add_sop_body_per_block(doc)
    p(doc, "{% if not loop.last %}")
    p(doc, "────────────────────────────────────────")
    p(doc, "{% endif %}")
    p(doc, "{% endfor %}")
    p(doc, "{% else %}")
    _add_sop_body_single(doc)
    p(doc, "{% endif %}")

    doc.save(str(out))
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
