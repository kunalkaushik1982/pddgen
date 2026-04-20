r"""
Generate docs/templates/flowlens-pdd-template.docx for docxtpl PDD export tests.

Multi-process sessions use ``pdd.multi_process`` and ``pdd.process_sections`` loops
inside the Word file (no runtime ``document.xml`` patching). Single-process
sessions use ``pdd.as_is_steps`` and the usual overview / diagram / rules blocks.

Also copies the output to ``test-assets/pdd-enterprise-multishot-template.docx``
so existing guide links and fixtures stay aligned.

Run from repo root: python docs/templates/build_flowlens_pdd_template.py
"""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document as open_document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def p(document: DocxDocument, text: str, *, style: str | None = None) -> None:
    if style:
        document.add_paragraph(text, style=style)
    else:
        document.add_paragraph(text)


def _add_single_process_body(doc: DocxDocument) -> None:
    """Sections 3–6 for one workflow area (matches docxtpl-template-guide lean structure)."""
    p(doc, "3. AS-IS Steps", style="Heading 1")
    p(doc, "{% for step in pdd.as_is_steps %}")
    p(doc, "{{ step.bullet_entry }}", style="List Bullet")
    p(doc, "{% if step.primary_screenshot_image %}")
    p(doc, "Primary screenshot:")
    p(doc, "{{ step.primary_screenshot_image }}")
    p(doc, "{% endif %}")
    p(doc, "{% endfor %}")

    p(doc, "")
    p(doc, "4. TO-BE Suggestions", style="Heading 1")
    p(doc, "{% for item in pdd.to_be_recommendations %}")
    p(doc, "{{ item.title }}: {{ item.recommendation }}", style="List Bullet")
    p(doc, "{% endfor %}")

    p(doc, "")
    p(doc, "5. Process Flow Diagram", style="Heading 1")
    p(doc, "{% if pdd.process_flow.detailed_image %}")
    p(doc, "{{ pdd.process_flow.detailed_image }}")
    p(doc, "{% elif pdd.process_flow.diagram_image %}")
    p(doc, "{{ pdd.process_flow.diagram_image }}")
    p(doc, "{% else %}")
    p(doc, "Diagram Source:")
    p(doc, "{{ pdd.process_flow.diagram_source }}")
    p(doc, "{% endif %}")

    p(doc, "")
    p(doc, "6. Business Rules and Notes", style="Heading 1")
    p(doc, "{% for rule in pdd.business_rules %}")
    p(doc, "{{ rule.text }} ({{ rule.inference_type }}, {{ rule.confidence }})", style="List Bullet")
    p(doc, "{% endfor %}")


def _add_multi_process_body(doc: DocxDocument) -> None:
    """Per-process sections then global TO-BE (replaces former OOXML injection)."""
    p(doc, "3. Process Sections", style="Heading 1")
    p(doc, "{% for section in pdd.process_sections %}")
    p(doc, "{{ section.title }}", style="Heading 2")
    p(doc, "{{ section.summary }}")
    p(doc, "AS-IS Steps", style="Heading 3")
    p(doc, "{% for step in section.steps %}")
    p(doc, "{{ step.bullet_entry }}", style="List Bullet")
    p(doc, "{% if step.primary_screenshot_image %}")
    p(doc, "Primary screenshot:")
    p(doc, "{{ step.primary_screenshot_image }}")
    p(doc, "{% endif %}")
    p(doc, "{% endfor %}")

    p(doc, "Process Flow Diagram", style="Heading 3")
    p(doc, "{% if section.diagram_image %}")
    p(doc, "{{ section.diagram_image }}")
    p(doc, "{% else %}")
    p(doc, "{{ section.diagram_source }}")
    p(doc, "{% endif %}")

    p(doc, "Business Rules and Notes", style="Heading 3")
    p(doc, "{% for rule in section.notes %}")
    p(doc, "{{ rule.text }}", style="List Bullet")
    p(doc, "{% endfor %}")

    p(doc, "{% if not loop.last %}")
    p(doc, "────────────────────────────────────────")
    p(doc, "{% endif %}")
    p(doc, "{% endfor %}")

    p(doc, "")
    p(doc, "4. TO-BE Suggestions", style="Heading 1")
    p(doc, "{% for item in pdd.to_be_recommendations %}")
    p(doc, "{{ item.title }}: {{ item.recommendation }}", style="List Bullet")
    p(doc, "{% endfor %}")


def main() -> None:
    root = Path(__file__).resolve().parents[2]
    out = Path(__file__).resolve().parent / "flowlens-pdd-template.docx"
    test_assets = root / "test-assets" / "pdd-enterprise-multishot-template.docx"

    doc = open_document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PROCESS DESIGN DOCUMENT")
    run.bold = True
    run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("{{ pdd.overview.process_name }}").italic = True

    p(doc, "")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "Document Owner: {{ pdd.overview.document_owner }}  |  "
        "Generated: {{ pdd.overview.generated_at }}  |  "
        "Session: {{ pdd.session_id }}  |  Status: {{ pdd.overview.document_status }}\n"
        "Diagram Type: {{ pdd.diagram_type }}  |  Steps: {{ pdd.step_count }}  |  Notes: {{ pdd.note_count }}"
    )

    p(doc, "")
    p(doc, "1. Document Overview", style="Heading 1")
    p(
        doc,
        "This document captures the AS-IS process observed during the discovery walkthrough and records draft TO-BE suggestions.",
    )

    p(doc, "")
    p(doc, "2. AS-IS Overview", style="Heading 1")
    p(doc, "{{ pdd.overview.process_summary }}")

    p(doc, "")
    p(doc, "{% if pdd.multi_process %}")
    _add_multi_process_body(doc)
    p(doc, "{% else %}")
    _add_single_process_body(doc)
    p(doc, "{% endif %}")

    doc.save(str(out))
    print(f"Wrote {out}")

    test_assets.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(out, test_assets)
    print(f"Copied to {test_assets}")


if __name__ == "__main__":
    main()
