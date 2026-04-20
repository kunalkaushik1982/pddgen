r"""
Generate docs/templates/flowlens-brd-template.docx for docxtpl BRD export tests.

Layout follows the standard BRD outline (index + detailed example style):
sections 1–20 plus 4.1 / 4.2 via ``brd.canonical_sections``, then supporting
loops for requirements, workflow evidence, and diagram.

When ``brd.multi_process`` is true (multiple workflow sections / process groups),
the same body repeats once per ``brd.process_document_blocks`` item; the cover
block (title, metadata, intro quote) stays single.

Run from repo root: python docs/templates/build_flowlens_brd_template.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document as open_document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def p(document: DocxDocument, text: str, *, style: str | None = None) -> None:
    if style:
        document.add_paragraph(text, style=style)
    else:
        document.add_paragraph(text)


def _add_brd_body_single(doc: DocxDocument) -> None:
    """Standard BRD sections + supporting appendix using top-level ``brd.*``."""
    p(doc, "")
    p(doc, "Standard BRD sections", style="Heading 1")

    p(doc, "{% for row in brd.canonical_sections %}")
    p(doc, "{{ row.ref }} {{ row.title }}", style="Heading 2")
    p(doc, "{{ row.body }}")
    p(doc, "{% endfor %}")

    p(doc, "")
    p(doc, "Document overview (metadata)", style="Heading 2")
    p(doc, "{{ brd.overview.process_name }}")
    p(doc, "{{ brd.overview.document_owner }} · {{ brd.overview.document_status }}")
    p(doc, "{{ brd.overview.process_summary }}")

    p(doc, "")
    p(doc, "Structured requirements (BR-###)", style="Heading 2")
    p(doc, "{% for req in brd.requirements %}")
    p(doc, "{{ req.id }} [{{ req.category }}]: {{ req.statement }} — {{ req.rationale }}")
    p(doc, "{% endfor %}")

    p(doc, "")
    p(doc, "Business rules and notes", style="Heading 2")
    p(doc, "{% for rule in brd.business_rules %}")
    p(doc, "{{ rule }}")
    p(doc, "{% endfor %}")

    p(doc, "")
    p(doc, "Workflow evidence (by section)", style="Heading 2")
    p(doc, "{% for sec in brd.workflow_sections %}")
    p(doc, "{{ sec.title }}: {{ sec.summary }} ({{ sec.step_count }} steps)")
    p(doc, "{% endfor %}")

    p(doc, "")
    p(doc, "Process flow (combined diagram)", style="Heading 2")
    p(doc, "{% if brd.process_flow.diagram_image %}{{ brd.process_flow.diagram_image }}{% endif %}")

    p(doc, "")
    p(doc, "Evidence summary", style="Heading 2")
    p(
        doc,
        "Sections: {{ brd.evidence_summary.workflow_sections }} · "
        "Steps: {{ brd.evidence_summary.observed_steps }} · "
        "Notes: {{ brd.evidence_summary.captured_notes }}",
    )


def _add_brd_body_per_block(doc: DocxDocument) -> None:
    """Same structure as single-process, but scoped to ``block`` inside the outer loop."""
    p(doc, "")
    p(doc, "{{ block.process_title }}", style="Heading 2")
    p(doc, "Standard BRD sections", style="Heading 1")

    p(doc, "{% for row in block.canonical_sections %}")
    p(doc, "{{ row.ref }} {{ row.title }}", style="Heading 2")
    p(doc, "{{ row.body }}")
    p(doc, "{% endfor %}")

    p(doc, "")
    p(doc, "Document overview (metadata)", style="Heading 2")
    p(doc, "{{ block.overview.process_name }}")
    p(doc, "{{ block.overview.document_owner }} · {{ block.overview.document_status }}")
    p(doc, "{{ block.overview.process_summary }}")

    p(doc, "")
    p(doc, "Structured requirements (BR-###)", style="Heading 2")
    p(doc, "{% for req in block.requirements %}")
    p(doc, "{{ req.id }} [{{ req.category }}]: {{ req.statement }} — {{ req.rationale }}")
    p(doc, "{% endfor %}")

    p(doc, "")
    p(doc, "Business rules and notes", style="Heading 2")
    p(doc, "{% for rule in block.business_rules %}")
    p(doc, "{{ rule }}")
    p(doc, "{% endfor %}")

    p(doc, "")
    p(doc, "Workflow evidence (by section)", style="Heading 2")
    p(doc, "{% for sec in block.workflow_sections %}")
    p(doc, "{{ sec.title }}: {{ sec.summary }} ({{ sec.step_count }} steps)")
    p(doc, "{% endfor %}")

    p(doc, "")
    p(doc, "Process flow (diagram for this process)", style="Heading 2")
    p(doc, "{% if block.process_flow.diagram_image %}{{ block.process_flow.diagram_image }}{% endif %}")

    p(doc, "")
    p(doc, "Evidence summary", style="Heading 2")
    p(
        doc,
        "Sections: {{ block.evidence_summary.workflow_sections }} · "
        "Steps: {{ block.evidence_summary.observed_steps }} · "
        "Notes: {{ block.evidence_summary.captured_notes }}",
    )


def main() -> None:
    out = Path(__file__).resolve().parent / "flowlens-brd-template.docx"
    doc = open_document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FlowLens Business Requirements Document")
    run.bold = True
    run.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("{{ brd.title }}").italic = True

    p(doc, "")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "Owner: {{ brd.owner_id }}  |  Session: {{ brd.session_id }}  |  Status: {{ brd.status }}\n"
        "Generated: {{ brd.generated_at }}  |  Diagram type: {{ brd.diagram_type }}"
    )

    p(doc, "")
    p(
        doc,
        "This document follows the Business Requirements Document (BRD) structure: executive summary "
        "through approval and sign-off. Section bodies are filled from the session evidence where "
        "available; placeholders indicate items to complete with stakeholders.",
        style="Intense Quote",
    )

    p(doc, "{% if brd.multi_process %}")
    p(doc, "{% for block in brd.process_document_blocks %}")
    _add_brd_body_per_block(doc)
    p(doc, "{% endfor %}")
    p(doc, "{% else %}")
    _add_brd_body_single(doc)
    p(doc, "{% endif %}")

    doc.save(str(out))
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
